#!/usr/bin/env python3
"""
Generate the Harari PCC Portal Project Proposal as a Word document.
Concise government-adoption proposal, ~10 pages, budget in ETB only.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# --- Harari palette ---
PURPLE = RGBColor(0x5B, 0x2A, 0x86)        # primary
PURPLE_DEEP = RGBColor(0x4A, 0x1F, 0x6E)
GOLD = RGBColor(0xD4, 0xA5, 0x37)          # accent
TERRACOTTA = RGBColor(0xB5, 0x47, 0x1A)
GREEN = RGBColor(0x2E, 0x7A, 0x5A)
INDIGO = RGBColor(0x1E, 0x3A, 0x5F)        # body text / ink
INK_MUTED = RGBColor(0x6B, 0x5B, 0x73)
CREAM = "FBF3E2"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

OUT = "/home/z/my-project/download/Harari_PCC_Portal_Proposal.docx"

doc = Document()

# --- Page setup: A4, standard margins ---
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# --- Default font (body): Times New Roman 12pt, line 1.3 ---
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.font.color.rgb = INDIGO
pf = style.paragraph_format
pf.line_spacing = 1.3
pf.space_before = Pt(0)
pf.space_after = Pt(6)

# East Asian font fallback (so the doc opens cleanly in WPS / MS Office)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:eastAsia"), "SimSun")
rfonts.set(qn("w:ascii"), "Times New Roman")
rfonts.set(qn("w:hAnsi"), "Times New Roman")

# --- Helpers ---

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def set_cell_borders(cell, color="D4A537", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)

def add_paragraph_border(paragraph, edge="bottom", color="D4A537", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = p_pr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        p_pr.append(pBdr)
    b = OxmlElement(f"w:{edge}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), size)
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), color)
    pBdr.append(b)

def shade_paragraph(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)

def h1(text, color=PURPLE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = color
    add_paragraph_border(p, edge="bottom", color="D4A537", size="6")
    return p

def h2(text, color=INDIGO):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = color
    return p

def h3(text, color=TERRACOTTA):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = color
    return p

def body(text, justify=True, indent=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = INDIGO
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    # clear default and add custom run
    p.text = ""
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = INDIGO
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = INK_MUTED
    return p

# =====================================================================
# COVER PAGE
# =====================================================================

# Top spacer
for _ in range(2):
    doc.add_paragraph()

# Top color stripe (using a 1x1 table with gold fill)
stripe = doc.add_table(rows=1, cols=1)
stripe.autofit = False
stripe.columns[0].width = Cm(15.5)
cell = stripe.cell(0, 0)
cell.width = Cm(15.5)
set_cell_bg(cell, "D4A537")
cell.paragraphs[0].text = ""
# Force thin height
tr = cell._tc.getparent()
trPr = tr.find(qn("w:trPr"))
if trPr is None:
    trPr = OxmlElement("w:trPr")
    tr.insert(0, trPr)
trHeight = OxmlElement("w:trHeight")
trHeight.set(qn("w:val"), "120")
trHeight.set(qn("w:hRule"), "exact")
trPr.append(trHeight)

# Spacer
doc.add_paragraph()

# English label
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HARARI PEOPLE REGIONAL STATE  ·  ETHIOPIA")
r.font.name = "Times New Roman"
r.font.size = Pt(11)
r.font.color.rgb = GOLD
r.font.bold = True
# Letter-spacing-ish via spaces is awkward in Word; keep simple

# Subtitle line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Trade, Industry & Tourism Development Bureau")
r.font.name = "Times New Roman"
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = INK_MUTED

# Big spacer
for _ in range(3):
    doc.add_paragraph()

# Main title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Harari PCC Portal")
r.font.name = "Times New Roman"
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = PURPLE

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run("Professional Competence Certificate Platform")
r.font.name = "Times New Roman"
r.font.size = Pt(18)
r.font.italic = True
r.font.color.rgb = TERRACOTTA

# Decorative divider line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph_border(p, edge="bottom", color="D4A537", size="12")
p.paragraph_format.space_after = Pt(20)

# Tagline
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run("A Proposal for Official Adoption and Regional Deployment")
r.font.name = "Times New Roman"
r.font.size = Pt(13)
r.font.color.rgb = INDIGO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("of a Digital Business Licensing Portal for the Harari Region")
r.font.name = "Times New Roman"
r.font.size = Pt(13)
r.font.color.rgb = INDIGO

# Big spacer
for _ in range(6):
    doc.add_paragraph()

# Meta block (as a small table for alignment)
meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.autofit = False
meta.columns[0].width = Cm(4.5)
meta.columns[1].width = Cm(9.0)
meta_data = [
    ("Submitted to:", "Harari Regional State Council"),
    ("Submitted by:", "Trade, Industry & Tourism Development Bureau"),
    ("Date:", "June 2026"),
    ("Version:", "1.0  ·  For Official Review"),
]
for i, (label, value) in enumerate(meta_data):
    c1 = meta.cell(i, 0)
    c2 = meta.cell(i, 1)
    c1.width = Cm(4.5)
    c2.width = Cm(9.0)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run(label)
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = INK_MUTED
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(value)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    r2.font.color.rgb = INDIGO
    # no borders
    for c in (c1, c2):
        tc_pr = c._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            tcBorders.append(b)
        tc_pr.append(tcBorders)

# Bottom color stripe
doc.add_paragraph()
stripe2 = doc.add_table(rows=1, cols=1)
stripe2.autofit = False
stripe2.columns[0].width = Cm(15.5)
cell2 = stripe2.cell(0, 0)
cell2.width = Cm(15.5)
set_cell_bg(cell2, "5B2A86")
cell2.paragraphs[0].text = ""
tr2 = cell2._tc.getparent()
trPr2 = tr2.find(qn("w:trPr"))
if trPr2 is None:
    trPr2 = OxmlElement("w:trPr")
    tr2.insert(0, trPr2)
trHeight2 = OxmlElement("w:trHeight")
trHeight2.set(qn("w:val"), "120")
trHeight2.set(qn("w:hRule"), "exact")
trPr2.append(trHeight2)

# Page break to body
doc.add_page_break()

# =====================================================================
# 1. EXECUTIVE SUMMARY
# =====================================================================
h1("1. Executive Summary")

body(
    "The Harari People Regional State, despite its cultural and commercial significance as the home of "
    "Harar Jugol \u2014 a UNESCO World Heritage Site \u2014 continues to rely on a paper-based, in-person "
    "business licensing process that is slow, opaque, and difficult to audit. Citizens travel to the "
    "Bureau office in Harar, queue for hours, submit paper documents by hand, and wait an average of "
    "five to fourteen business days for a decision. There is no public registry to verify issued "
    "certificates, leaving room for forgery and undermining trust in the regional licensing system."
)

body(
    "This proposal presents the Harari PCC Portal \u2014 a complete, working digital platform that moves "
    "the entire Professional Competence Certificate (PCC) workflow online. Citizens register, submit "
    "documents, complete a competence assessment, and receive a verifiable certificate without visiting "
    "a physical office. Regional officers review applications through a dedicated console with full "
    "audit logging. The platform is built on modern, secure, scalable technology and is already "
    "functional as a working prototype."
)

body(
    "We request the Regional Council's formal adoption of this platform as the official digital channel "
    "for PCC issuance, along with a budget allocation of ETB 2,460,000 for production deployment, "
    "capacity building, and the first twelve months of operation. The expected outcomes are: a "
    "reduction in average processing time from 5\u201314 days to under 24 hours; full visibility of "
    "application status for every citizen; and a public certificate verification endpoint that "
    "eliminates forgery risk. The platform aligns with federal digital transformation priorities and "
    "with the Harari Region's broader goal of becoming a model small-region administration in Ethiopia."
)

# Highlight box
hl = doc.add_paragraph()
hl.paragraph_format.space_before = Pt(8)
hl.paragraph_format.space_after = Pt(8)
shade_paragraph(hl, "FBF3E2")
add_paragraph_border(hl, edge="left", color="D4A537", size="18")
add_paragraph_border(hl, edge="top", color="D4A537", size="6")
add_paragraph_border(hl, edge="bottom", color="D4A537", size="6")
add_paragraph_border(hl, edge="right", color="D4A537", size="6")
hl.paragraph_format.left_indent = Cm(0.3)
hl.paragraph_format.right_indent = Cm(0.3)
r = hl.add_run("Key Ask:  ")
r.font.bold = True
r.font.color.rgb = PURPLE
r.font.size = Pt(11)
r2 = hl.add_run(
    "Formal adoption of the Harari PCC Portal as the official digital PCC channel, plus a "
    "twelve-month implementation and operating budget of ETB 2,460,000."
)
r2.font.size = Pt(11)
r2.font.color.rgb = INDIGO

# =====================================================================
# 2. BACKGROUND & PROBLEM ANALYSIS
# =====================================================================
h1("2. Background and Problem Analysis")

h2("2.1  Current State of Business Licensing in the Harari Region")

body(
    "Under the existing arrangement, any citizen of the Harari Region who wishes to open a business "
    "must obtain a Professional Competence Certificate (PCC) from the Trade, Industry & Tourism "
    "Development Bureau before applying for a trade licence. The PCC process is designed to verify "
    "that the applicant understands basic Ethiopian business law, taxation, labour regulations, and "
    "regional requirements, and that they hold the educational or professional qualifications "
    "appropriate to their intended business sector."
)

body(
    "In practice, every step of this process is paper-based. The applicant obtains a paper form from "
    "the Bureau office in Harar, completes it by hand, attaches photocopies of supporting documents "
    "(national ID, educational certificates, business plan, lease agreement), and submits the folder "
    "in person. The folder is then routed manually between the front desk, a reviewing officer, and "
    "the head of the licensing division. The applicant has no way to know where their file is in the "
    "queue. When the decision is ready, the applicant must return to the Bureau to collect the "
    "certificate in person."
)

h2("2.2  Specific Pain Points")

bullet(
    "Processing time of 5\u201314 business days per application, with peaks of three weeks during "
    "high-demand periods such as the run-up to trade-fair season."
)
bullet(
    "No status visibility \u2014 applicants cannot track their application and must call or visit the "
    "Bureau to enquire, consuming staff time on both sides."
)
bullet(
    "No public verification \u2014 third parties (banks, landlords, partner businesses) cannot "
    "validate an applicant's certificate without phoning the Bureau, creating forgery risk and "
    "friction in commercial transactions."
)
bullet(
    "Lost or misrouted files \u2014 paper folders are occasionally misplaced during internal routing, "
    "forcing applicants to resubmit documents and inflating the actual processing time further."
)
bullet(
    "Limited service reach \u2014 citizens living in woredas outside Harar must travel to the regional "
    "capital for every step of the process, imposing a hidden cost on rural entrepreneurs."
)
bullet(
    "No management analytics \u2014 the Bureau leadership has no real-time view of pending "
    "applications, sectoral distribution, or reviewer workload, making resource planning reactive."
)

h2("2.3  Policy and Strategic Context")

body(
    "The federal government of Ethiopia has, since 2022, made digital transformation a national "
    "priority under the \"Digital Ethiopia 2025\" strategy. Several regional states \u2014 including "
    "Oromia, Amhara, and the Sidama Region \u2014 have begun deploying online portals for selected "
    "trade and land services. The Harari Region, despite being the smallest regional state by area, "
    "has lagged behind in this regard, in part because no tailored solution existed that matched the "
    "Region's specific Bureau workflow and cultural context."
)

body(
    "Adopting the Harari PCC Portal directly addresses this gap. It is built specifically for the "
    "Harari Region \u2014 not a generic off-the-shelf product \u2014 and aligns with the regional "
    "Bureau's actual decision-making workflow, document requirements, and cultural identity."
)

# =====================================================================
# 3. PROJECT OBJECTIVES AND EXPECTED OUTCOMES
# =====================================================================
h1("3. Project Objectives and Expected Outcomes")

h2("3.1  Primary Objectives")

body(
    "The project has four primary objectives, each tied to a measurable outcome:"
)

bullet(
    "Digitise the end-to-end PCC workflow \u2014 from registration to certificate issuance \u2014 so "
    "that no physical office visit is required for any step of an application."
)
bullet(
    "Enforce competence verification through a structured 17-question assessment covering Ethiopian "
    "business law, taxation, labour law, consumer protection, and Harari regional regulations."
)
bullet(
    "Provide a public, no-authentication certificate verification endpoint so that any third party "
    "can validate an issued certificate by its unique number."
)
bullet(
    "Give Bureau leadership a real-time dashboard of application volume, status distribution, "
    "sectoral breakdown, and reviewer activity for operational decision-making."
)

h2("3.2  Expected Outcomes and Indicators")

# Outcomes table
outcomes = [
    ("Outcome", "Indicator", "Baseline", "Target (12 mo)"),
    ("Faster processing",
     "Average time from submission to decision",
     "5\u201314 days",
     "Under 24 hours"),
    ("Citizen transparency",
     "Share of applicants who can check status online",
     "0%",
     "100%"),
    ("Fraud reduction",
     "Public verification endpoint availability",
     "None",
     "Live, no-auth endpoint"),
    ("Geographic reach",
     "Applications submitted from outside Harar city",
     "Unknown",
     "Tracked and reported quarterly"),
    ("Operational visibility",
     "Bureau dashboard with live queue + analytics",
     "None",
     "Deployed and used daily"),
]

t = doc.add_table(rows=len(outcomes), cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
widths = [Cm(3.6), Cm(5.6), Cm(2.7), Cm(3.4)]
for row_idx, row_data in enumerate(outcomes):
    row = t.rows[row_idx]
    for col_idx, value in enumerate(row_data):
        cell = row.cells[col_idx]
        cell.width = widths[col_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        r.font.name = "Times New Roman"
        if row_idx == 0:
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = WHITE
            set_cell_bg(cell, "5B2A86")
        else:
            r.font.size = Pt(11)
            r.font.color.rgb = INDIGO
            if row_idx % 2 == 0:
                set_cell_bg(cell, "FBF3E2")
            else:
                set_cell_bg(cell, "FFFFFF")
        set_cell_borders(cell, color="D4A537", size="4")
        if row_idx == 0:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT

caption("Table 1. Expected outcomes and measurable indicators.")

# =====================================================================
# 4. SOLUTION DESCRIPTION
# =====================================================================
h1("4. Solution Description")

h2("4.1  Platform Overview")

body(
    "The Harari PCC Portal is a web-based platform accessible from any modern browser on desktop or "
    "mobile. It serves three user roles with strictly enforced, server-side access control: "
    "applicants, reviewers (Bureau officers), and administrators. Each role sees only the "
    "functionality and data appropriate to it."
)

body(
    "The platform implements a five-step application wizard that guides the applicant through "
    "personal information, business information, document upload, competence assessment, and a final "
    "review-and-submit step. A progress indicator at the top of the wizard shows the applicant "
    "exactly where they are in the process. Drafts are saved automatically, so an applicant can "
    "pause at any step and return later without losing their work."
)

h2("4.2  Key Functional Capabilities")

h3("For Applicants")
bullet("Self-service registration with national ID, contact information, and regional address details.")
bullet("Five-step guided application wizard with autosave and progress tracking.")
bullet("Document upload supporting 19 document types aligned to the Ethiopian education system, including Grade 8, Grade 10 (EGECE), Grade 12 (matric), TVET certificates, diplomas, advanced diplomas, bachelor's, master's, and doctoral degrees, plus professional certifications and supporting business documents.")
bullet("Randomised 10-question competence assessment drawn from a 17-question bank, with a 70% pass mark and unlimited retries. Each incorrect answer is followed by an explanation so the applicant learns from every attempt.")
bullet("Real-time status tracking (Draft \u2192 Submitted \u2192 Under Review \u2192 Approved / Rejected \u2192 Certificate Issued) with a full audit-trail timeline visible per application.")
bullet("One-click certificate viewing and printing, with the option to save the certificate as a PDF.")

h3("For Reviewers and Bureau Officers")
bullet("Application queue with full-text search and status filters (Submitted, Under Review, Issued, Rejected).")
bullet("One-click workflow: Claim \u2192 Approve (auto-issues certificate) or Reject (with mandatory written reason).")
bullet("Educational qualifications summary card per applicant, displaying all verified credentials at a glance.")
bullet("Dashboard analytics with counts by status, certificates issued by sector, and a live recent-activity feed.")
bullet("Full audit log \u2014 every action (login, claim, approve, reject, issue) is recorded with user, timestamp, and IP address, supporting accountability and after-the-fact investigation.")

h3("For the Public")
bullet("A public certificate verification endpoint. Any third party \u2014 a bank, a landlord, a partner business, a customer \u2014 can validate a certificate by entering its unique number (e.g. HRS-PCC-CERT-2026-0001) and receiving an instant valid/invalid response with the certificate holder's name, business name, and validity period.")

h2("4.3  Cultural and Regional Alignment")

body(
    "Unlike generic licensing software, the Harari PCC Portal is designed to feel native to the "
    "Region. The visual identity draws on Harari cultural motifs \u2014 the eight-pointed star, the "
    "Harar Jugol gate, and a colour palette of royal purple, gold, terracotta, cream, and Islamic "
    "green inspired by the historic walled city. The institutional emblem combines traditional "
    "Ethiopian symbolic elements (balance scale for justice, sun for hope, castle for heritage, gear "
    "for industry) reimagined in the regional palette. This cultural rooting matters: citizens are "
    "more likely to trust and adopt a service that visibly belongs to their own Region rather than "
    "a generic national template."
)

# =====================================================================
# 5. TECHNICAL ARCHITECTURE
# =====================================================================
h1("5. Technical Architecture")

h2("5.1  Technology Stack")

body(
    "The platform is built on a modern, industry-standard technology stack that is mature, "
    "well-documented, and widely supported by Ethiopian and international developers. This ensures "
    "the Region is not locked into a single vendor and can hire local talent to maintain and extend "
    "the system."
)

stack = [
    ("Layer", "Technology", "Role"),
    ("Frontend", "Next.js 16 + TypeScript + Tailwind CSS + shadcn/ui",
     "Responsive web interface, mobile-first design"),
    ("Backend", "Next.js API routes (Node.js 20+)",
     "REST API with 12 endpoint groups"),
    ("Database", "Prisma ORM + PostgreSQL (production)",
     "Relational storage with type-safe access"),
    ("Authentication", "bcryptjs password hashing + JWT sessions in httpOnly cookies",
     "Secure, stateless, scalable auth"),
    ("File storage", "S3-compatible object storage (e.g., MinIO, AWS S3)",
     "Encrypted document storage with signed URLs"),
    ("Security", "Role-based access control enforced server-side",
     "Applicant / Reviewer / Admin permissions"),
]

t2 = doc.add_table(rows=len(stack), cols=3)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
t2.autofit = False
w2 = [Cm(3.0), Cm(6.5), Cm(5.8)]
for row_idx, row_data in enumerate(stack):
    row = t2.rows[row_idx]
    for col_idx, value in enumerate(row_data):
        cell = row.cells[col_idx]
        cell.width = w2[col_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        r.font.name = "Times New Roman"
        if row_idx == 0:
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = WHITE
            set_cell_bg(cell, "5B2A86")
        else:
            r.font.size = Pt(11)
            r.font.color.rgb = INDIGO
            if row_idx % 2 == 0:
                set_cell_bg(cell, "FBF3E2")
            else:
                set_cell_bg(cell, "FFFFFF")
        set_cell_borders(cell, color="D4A537", size="4")

caption("Table 2. Production technology stack.")

h2("5.2  Data Model and Security")

body(
    "The platform stores seven core entities: user profiles, applications, documents, certificates, "
    "audit logs, notifications, and a sequence counter for reference-number generation. All "
    "sensitive fields \u2014 passwords, document files, audit logs \u2014 are protected at the "
    "application layer (bcrypt hashing, signed-URL access, server-side role checks) and at the "
    "database layer (PostgreSQL row-level security policies can be enabled for defence in depth)."
)

body(
    "Document files (PDFs, images of certificates, lease agreements) are stored in an S3-compatible "
    "object store rather than in the database itself, which keeps the database small and fast and "
    "allows individual documents to be expired, backed up, or migrated independently. Access to "
    "documents is mediated by short-lived signed URLs that expire within sixty seconds, so even if a "
    "URL is leaked it cannot be reused."
)

# =====================================================================
# 6. IMPLEMENTATION PLAN
# =====================================================================
h1("6. Implementation Plan")

h2("6.1  Phased Approach")

body(
    "We propose a four-phase implementation over twelve months. Each phase concludes with a concrete "
    "deliverable that the Bureau can review and approve before the next phase begins."
)

phases = [
    ("Phase", "Duration", "Key Activities", "Deliverable"),
    ("1. Production deployment",
     "Months 1\u20132",
     "Provision production infrastructure (database, storage, hosting). Migrate from SQLite to PostgreSQL. Configure backups, monitoring, and a custom domain (pcc.harariregion.gov.et).",
     "Live production environment"),
    ("2. Pilot rollout",
     "Months 3\u20134",
     "Onboard the first cohort of 50 applicants and 3 reviewers from the Bureau. Run parallel paper + digital processing. Collect feedback weekly.",
     "Pilot evaluation report"),
    ("3. Full rollout & training",
     "Months 5\u20138",
     "Train all Bureau licensing staff. Open the portal to the general public. Decommission paper-only intake. Run public awareness campaign.",
     "Public launch + trained staff"),
    ("4. Stabilisation & handover",
     "Months 9\u201312",
     "Bug fixes, performance tuning, first analytics review. Hand over operational responsibility to the Bureau IT unit. Prepare roadmap for year-2 enhancements.",
     "Operational handover"),
]

t3 = doc.add_table(rows=len(phases), cols=4)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
t3.autofit = False
w3 = [Cm(3.4), Cm(2.2), Cm(7.0), Cm(3.0)]
for row_idx, row_data in enumerate(phases):
    row = t3.rows[row_idx]
    for col_idx, value in enumerate(row_data):
        cell = row.cells[col_idx]
        cell.width = w3[col_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        r.font.name = "Times New Roman"
        if row_idx == 0:
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = WHITE
            set_cell_bg(cell, "5B2A86")
        else:
            r.font.size = Pt(11)
            r.font.color.rgb = INDIGO
            if row_idx % 2 == 0:
                set_cell_bg(cell, "FBF3E2")
            else:
                set_cell_bg(cell, "FFFFFF")
        set_cell_borders(cell, color="D4A537", size="4")

caption("Table 3. Four-phase implementation plan over twelve months.")

h2("6.2  Key Milestones")

bullet("Month 2: Production environment live at pcc.harariregion.gov.et.")
bullet("Month 4: Pilot evaluation report delivered to the Bureau leadership.")
bullet("Month 5: All Bureau licensing staff trained and certified on the platform.")
bullet("Month 6: Public launch; paper-only intake decommissioned.")
bullet("Month 12: Operational handover to the Bureau IT unit; year-2 roadmap agreed.")

# =====================================================================
# 7. BUDGET
# =====================================================================
h1("7. Budget")

body(
    "The total requested budget for the twelve-month implementation and first-year operation is "
    "ETB 2,460,000. The budget is divided into four categories: one-time setup, software "
    "development and customisation, capacity building and training, and recurring operating costs. "
    "All figures are in Ethiopian Birr and are estimates based on typical Ethiopian IT project "
    "costing; final figures will be confirmed during procurement."
)

budget = [
    ("Category", "Line item", "ETB"),
    ("One-time setup",
     "Domain registration (pcc.harariregion.gov.et) \u2014 3 years",
     "15,000"),
    ("", "SSL certificate \u2014 3 years", "30,000"),
    ("", "Production hosting (cloud VPS, 4 vCPU / 8 GB) \u2014 12 mo", "180,000"),
    ("", "PostgreSQL managed database \u2014 12 mo", "144,000"),
    ("", "S3-compatible object storage (50 GB) \u2014 12 mo", "60,000"),
    ("", "Backup and disaster recovery setup", "75,000"),
    ("Development",
     "Production hardening, security audit, PostgreSQL migration",
     "350,000"),
    ("", "Integration with ERCA TIN verification API", "180,000"),
    ("", "Amharic / Arabic localisation (UI + assessment)", "120,000"),
    ("", "Mobile-responsive optimisation and PWA support", "150,000"),
    ("", "Bug fixing and feature polishing (months 9\u201312)", "200,000"),
    ("Capacity building",
     "Training of 10 Bureau staff (3 days @ ETB 8,000/day)",
     "240,000"),
    ("", "Training materials, manuals, video tutorials", "60,000"),
    ("", "Public awareness campaign (radio, social, flyers)", "150,000"),
    ("Operations",
     "System administrator (part-time, 12 mo)", "300,000"),
    ("", "Monitoring, security patching, incident response", "120,000"),
    ("", "Contingency (10%)", "186,000"),
    ("", "TOTAL", "2,460,000"),
]

t4 = doc.add_table(rows=len(budget), cols=3)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
t4.autofit = False
w4 = [Cm(3.4), Cm(9.0), Cm(3.0)]
for row_idx, row_data in enumerate(budget):
    row = t4.rows[row_idx]
    for col_idx, value in enumerate(row_data):
        cell = row.cells[col_idx]
        cell.width = w4[col_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if col_idx == 2 and row_idx > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(value)
        r.font.name = "Times New Roman"
        if row_idx == 0:
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = WHITE
            set_cell_bg(cell, "5B2A86")
        elif row_idx == len(budget) - 1:
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = WHITE
            set_cell_bg(cell, "B5471A")
        else:
            r.font.size = Pt(11)
            r.font.color.rgb = INDIGO
            if row_idx % 2 == 0:
                set_cell_bg(cell, "FBF3E2")
            else:
                set_cell_bg(cell, "FFFFFF")
        set_cell_borders(cell, color="D4A537", size="4")

caption("Table 4. Detailed project budget in Ethiopian Birr (ETB).")

# =====================================================================
# 8. RISK ANALYSIS AND MITIGATION
# =====================================================================
h1("8. Risk Analysis and Mitigation")

body(
    "We have identified five principal risks to successful delivery and adoption. None of them is "
    "unusual for a government digital-service project, and each has a concrete mitigation plan."
)

risks = [
    ("Risk", "Likelihood", "Impact", "Mitigation"),
    ("Low digital literacy among rural applicants",
     "Medium", "Medium",
     "Provide a printable step-by-step guide in Amharic; offer assisted-application windows at the Bureau office during the first six months."),
    ("Staff resistance to new workflow",
     "Medium", "High",
     "Involve Bureau staff in pilot feedback; run three-day hands-on training; designate a staff champion in each division."),
    ("Internet connectivity outages",
     "Medium", "Medium",
     "The portal is designed to gracefully resume in-progress applications; offline data entry can be done at the Bureau office using the staff console."),
    ("Data breach or document leakage",
     "Low", "High",
     "All documents stored in encrypted object storage; access mediated by short-lived signed URLs; full audit log of every access; quarterly security review."),
    ("Scope creep (new feature requests during build)",
     "High", "Medium",
     "Strict change-control process: any new feature requested after Month 2 is logged and scheduled for Year 2 unless it blocks launch."),
]

t5 = doc.add_table(rows=len(risks), cols=4)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
t5.autofit = False
w5 = [Cm(4.5), Cm(2.0), Cm(1.8), Cm(7.0)]
for row_idx, row_data in enumerate(risks):
    row = t5.rows[row_idx]
    for col_idx, value in enumerate(row_data):
        cell = row.cells[col_idx]
        cell.width = w5[col_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        r.font.name = "Times New Roman"
        if row_idx == 0:
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = WHITE
            set_cell_bg(cell, "5B2A86")
        else:
            r.font.size = Pt(11)
            r.font.color.rgb = INDIGO
            if row_idx % 2 == 0:
                set_cell_bg(cell, "FBF3E2")
            else:
                set_cell_bg(cell, "FFFFFF")
        set_cell_borders(cell, color="D4A537", size="4")

caption("Table 5. Risk register with mitigation actions.")

# =====================================================================
# 9. SUSTAINABILITY
# =====================================================================
h1("9. Sustainability")

body(
    "Beyond the first twelve months, the platform is designed to be sustainable in three respects: "
    "technical, financial, and institutional."
)

h2("9.1  Technical Sustainability")

body(
    "The technology stack is built on widely-used open-source frameworks (Next.js, PostgreSQL, "
    "Prisma). This means the Bureau is not dependent on any single vendor for maintenance or "
    "enhancements \u2014 any competent Ethiopian software developer can be onboarded to support the "
    "system. The codebase is TypeScript throughout, which reduces the cost of future changes by "
    "catching errors at compile time. Documentation \u2014 including a comprehensive README, API "
    "reference, and operational runbook \u2014 will be delivered as part of the handover."
)

h2("9.2  Financial Sustainability")

body(
    "Recurring annual operating costs after Year 1 are estimated at ETB 420,000 (hosting, database, "
    "storage, domain, SSL, part-time system administrator). These costs can be absorbed into the "
    "Bureau's regular IT operating budget. The Region may also, in future, introduce a small "
    "digital-service fee per application (as several other regions have done) to fully recover "
    "operating costs from applicants who choose to use the digital channel."
)

h2("9.3  Institutional Sustainability")

body(
    "Operational responsibility for the platform will be transferred to the Bureau IT unit by Month "
    "12. The handover will include: source code, documentation, operational runbooks, monitoring "
    "dashboards, and a one-month shadow-support period during which the development team remains "
    "available for escalation. The Bureau IT unit will be trained to handle routine operations "
    "(user management, application review support, certificate verification, backup verification) "
    "without external assistance."
)

# =====================================================================
# 10. CONCLUSION AND REQUESTED ACTION
# =====================================================================
h1("10. Conclusion and Requested Action")

body(
    "The Harari PCC Portal is a complete, working, region-specific digital platform that addresses "
    "a real, documented pain point in the daily experience of Harari entrepreneurs. It is not a "
    "concept or a mockup \u2014 it is a functioning prototype that can be deployed to production "
    "within two months of approval. It aligns with federal digital-transformation priorities, "
    "respects the Region's cultural identity, and is built on a maintainable open-source stack."
)

body(
    "We respectfully request the Regional Council to:"
)

bullet("Formally adopt the Harari PCC Portal as the official digital channel for Professional Competence Certificate issuance in the Harari Region.")
bullet("Approve a twelve-month implementation and operating budget of ETB 2,460,000, drawn from the Bureau's innovation and digital-transformation allocation.")
bullet("Designate a project sponsor at the level of Bureau Head, with authority to coordinate across IT, licensing, and communications functions.")
bullet("Authorize a public launch by Month 6 of the implementation timeline, accompanied by the planned awareness campaign.")

body(
    "We are available at the Council's convenience to present the working prototype in a live "
    "demonstration, to answer technical or operational questions, and to refine the proposal based "
    "on the Council's guidance.",
    indent=False
)

# Signature block
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Submitted by:")
r.font.name = "Times New Roman"
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = INK_MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Trade, Industry & Tourism Development Bureau")
r.font.name = "Times New Roman"
r.font.size = Pt(12)
r.font.bold = True
r.font.color.rgb = PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Harari People Regional State  ·  June 2026")
r.font.name = "Times New Roman"
r.font.size = Pt(11)
r.font.color.rgb = INDIGO

# Save
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT)} bytes")
